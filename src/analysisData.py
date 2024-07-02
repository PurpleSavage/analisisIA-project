import pandas as pd
from connection import ConnectionDb
from docx import Document
import os
import shutil

class Data:
    def get_list_by_employee(self,id_empleado):
        with ConnectionDb() as connection:
            if connection:
                cursor = connection.cursor()
                
        
                # Consulta para obtener el historial del empleado
                query = """
                    SELECT idHistorial, idEmpleado, Fecha, Hora
                    FROM historial
                    WHERE idEmpleado = %s
                """
                cursor.execute(query, (id_empleado,))
                historial = cursor.fetchall()
                
                return historial
    def get_historial(self):
        with ConnectionDb() as connection:
            if connection:
                cursor =connection.cursor()
                    
                query = """
                    SELECT 
                            h.idEmpleado, 
                            e.nombre, 
                            e.apellido, 
                            e.dni, 
                            COUNT(h.idEmpleado) as cantidad_asistencias
                        FROM historial h
                        JOIN Empleado e ON h.idEmpleado = e.idEmpleado
                        GROUP BY h.idEmpleado, e.nombre, e.apellido, e.dni
                    """
                cursor.execute(query)
                historial=cursor.fetchall()
                return historial
    def get_percentage_assist(self,list):
        newlist = []
        total_dias=5
        for empleado in list:
            dict_empleado= {
                "id":empleado[0],
                "nombre":empleado[1],
                "apellido":empleado[2],
                "dni":empleado[3],
                "cantidad_asistencias": empleado[4],
                "porcentaje_asistencias": (empleado[4] / total_dias) * 100
            }
            newlist.append(dict_empleado)         
        return newlist

    def create_word(self, data_list, texto, nombre_archivo):
        if not data_list:
            print("La lista de datos está vacía.")
            return

        # Crear el DataFrame y asignar nombres a las columnas
        df = pd.DataFrame(data_list)
        df.columns = ['id', 'nombre', 'apellido', 'dni', 'n° asistencias', '"%"-asistencias']

        # Crear el documento de Word
        doc = Document()
        
        # Añadir el texto proporcionado
        doc.add_paragraph(texto)
        doc.add_paragraph("")  

        # Añadir una tabla al documento
        t = doc.add_table(df.shape[0] + 1, df.shape[1])
        t.style = 'Table Grid'
        # Añadir las cabeceras de la tabla
        for j in range(df.shape[-1]):
            t.cell(0, j).text = df.columns[j]

        # Añadir el resto de los datos del DataFrame a la tabla
        for i in range(df.shape[0]):
            for j in range(df.shape[-1]):
                t.cell(i + 1, j).text = str(df.values[i, j])

        # Definir la ruta del escritorio y del archivo de destino
        ruta_escritorio = os.path.join(os.path.expanduser('~'), 'Desktop')
        ruta_destino = os.path.join(ruta_escritorio, 'mi_carpeta', f'{nombre_archivo}.docx')

        try:
            # Verificar si la carpeta de destino existe, si no, crearla
            if not os.path.exists(os.path.dirname(ruta_destino)):
                os.makedirs(os.path.dirname(ruta_destino))
            
            # Guardar el documento de Word en la ruta de destino
            doc.save(ruta_destino)
            print(f'Documento guardado exitosamente en {ruta_destino}')
        except Exception as e:
            print(f'Error al guardar el documento: {e}')
 

    